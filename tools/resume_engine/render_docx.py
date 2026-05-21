"""
Render a markdown document (resume, cover letter, etc.) into a clean .docx file.

Usage:
    python tools/resume_engine/render_docx.py [path/to/document.md]

If no path is given, defaults to Resume - Master.md.
Outputs: <basename>.docx alongside the source file + a copy to ~/Downloads.
"""

import os
import re
import sys
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
FONT_NAME = "Calibri"
FONT_SIZE_H1 = Pt(15)
FONT_SIZE_H2 = Pt(13)
FONT_SIZE_H3 = Pt(11)
FONT_SIZE_H4 = Pt(10)
FONT_SIZE_BODY = Pt(10)

COLOR_H2 = RGBColor(0x34, 0x98, 0xDB)  # accent blue

MARGIN_TOP = Inches(0.25)
MARGIN_BOTTOM = Inches(0.25)
MARGIN_LEFT = Inches(0.5)
MARGIN_RIGHT = Inches(0.5)

DEFAULT_SOURCE = Path(__file__).resolve().parent.parent.parent / "Vault" / "3. Operations & Wealth" / "3.1. Career Strategy & Revenue" / "3.1.3. Professional Portfolio & Evidence" / "Resumes" / "Resume - Master.md"


# ---------------------------------------------------------------------------
# Inline formatting helpers
# ---------------------------------------------------------------------------
def _set_run_font(run, size=FONT_SIZE_BODY, bold=False, italic=False, color=None):
    """Apply font properties to a single run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _set_paragraph_spacing(paragraph, before=Pt(0), after=Pt(0), line_spacing=1.0):
    """Set spacing for a paragraph."""
    pf = paragraph.paragraph_format
    pf.space_before = before
    pf.space_after = after
    pf.line_spacing = line_spacing


def _add_hyperlink(paragraph, url, text, font_size=FONT_SIZE_BODY):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})

    run_elem = paragraph._element.makeelement(qn("w:r"), {})
    rPr = paragraph._element.makeelement(qn("w:rPr"), {})

    # Font name
    rFonts = paragraph._element.makeelement(qn("w:rFonts"), {qn("w:ascii"): FONT_NAME, qn("w:hAnsi"): FONT_NAME})
    rPr.append(rFonts)

    # Font size
    sz = paragraph._element.makeelement(qn("w:sz"), {qn("w:val"): str(int(font_size.pt * 2))})
    rPr.append(sz)

    # Blue + underline
    color_elem = paragraph._element.makeelement(qn("w:color"), {qn("w:val"): "3498DB"})
    rPr.append(color_elem)
    u_elem = paragraph._element.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rPr.append(u_elem)

    run_elem.append(rPr)

    t_elem = paragraph._element.makeelement(qn("w:t"), {})
    t_elem.text = text
    t_elem.set(qn("xml:space"), "preserve")
    run_elem.append(t_elem)

    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


def _add_formatted_text(paragraph, text, font_size=FONT_SIZE_BODY, base_bold=False, base_italic=False):
    """
    Parse inline markdown formatting and add runs to a paragraph.
    Handles: **bold**, *italic*, [text](url)
    """
    # Pattern to match **bold**, *italic*, or [text](url)
    pattern = re.compile(
        r'\*\*(.+?)\*\*'        # **bold**
        r'|\*(.+?)\*'           # *italic*
        r'|\[([^\]]+)\]\(([^)]+)\)'  # [text](url)
    )

    pos = 0
    for match in pattern.finditer(text):
        # Add text before this match
        before = text[pos:match.start()]
        if before:
            run = paragraph.add_run(before)
            _set_run_font(run, size=font_size, bold=base_bold, italic=base_italic)

        if match.group(1):  # **bold**
            run = paragraph.add_run(match.group(1))
            _set_run_font(run, size=font_size, bold=True, italic=base_italic)
        elif match.group(2):  # *italic*
            run = paragraph.add_run(match.group(2))
            _set_run_font(run, size=font_size, bold=base_bold, italic=True)
        elif match.group(3):  # [text](url)
            _add_hyperlink(paragraph, match.group(4), match.group(3), font_size=font_size)

        pos = match.end()

    # Add remaining text
    remaining = text[pos:]
    if remaining:
        run = paragraph.add_run(remaining)
        _set_run_font(run, size=font_size, bold=base_bold, italic=base_italic)


def _add_h2_bottom_border(paragraph):
    """Add a thin blue bottom border to an H2 paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "3498DB",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Markdown line parser
# ---------------------------------------------------------------------------
def _parse_heading(line):
    """Return (level, content) if line is a heading, else None."""
    m = re.match(r'^(#{1,6})\s+(.+)$', line)
    if m:
        return len(m.group(1)), m.group(2)
    return None


def _parse_bullet(line):
    """Return bullet content if line starts with '- ', else None."""
    m = re.match(r'^-\s+(.+)$', line)
    return m.group(1) if m else None


# ---------------------------------------------------------------------------
# Main renderer
# ---------------------------------------------------------------------------
def render_docx(source_path=None):
    source = Path(source_path) if source_path else DEFAULT_SOURCE
    if not source.exists():
        print(f"Error: Document not found at {source}")
        return

    md = source.read_text(encoding="utf-8")

    # Strip YAML frontmatter
    md = re.sub(r'^---.*?---\r?\n', '', md, count=1, flags=re.DOTALL)
    # Strip Obsidian navigation links ("Back to:" lines with wiki-links)
    md = re.sub(r'^Back to:.*\r?\n?', '', md, count=1, flags=re.MULTILINE)
    lines = md.strip().splitlines()

    # Derive output name from source filename
    base_name = source.stem  # e.g. "Resume - PowerMarket" or "Cover Letter - PowerMarket"
    output_dir = source.parent

    # Create document
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)

    # Set default font via document style
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Process lines
    for line in lines:
        line = line.rstrip()

        # Skip blank lines
        if not line.strip():
            continue

        # Check for heading
        heading = _parse_heading(line)
        if heading:
            level, content = heading

            if level == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _set_paragraph_spacing(p, before=Pt(0), after=Pt(0))
                _add_formatted_text(p, content, font_size=FONT_SIZE_H1, base_bold=True)
                continue

            if level == 2:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=Pt(6), after=Pt(2))
                _add_formatted_text(p, content, font_size=FONT_SIZE_H2, base_bold=True)
                _add_h2_bottom_border(p)
                continue

            if level == 3:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=Pt(3), after=Pt(1))
                _add_formatted_text(p, content, font_size=FONT_SIZE_H3, base_bold=True)
                continue

            if level == 4:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=Pt(2), after=Pt(1))
                _add_formatted_text(p, content, font_size=FONT_SIZE_H4, base_bold=False)
                continue

        # Check for bullet
        bullet = _parse_bullet(line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            _set_paragraph_spacing(p, before=Pt(0), after=Pt(1), line_spacing=1.0)
            _add_formatted_text(p, bullet, font_size=FONT_SIZE_BODY)
            continue

        # Regular paragraph (e.g., contact info line)
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=Pt(0), after=Pt(1))
        _add_formatted_text(p, line, font_size=FONT_SIZE_BODY)

    # Save outputs
    output_path = output_dir / f"{base_name}.docx"
    downloads_path = Path(os.environ.get("USERPROFILE", "~")) / "Downloads" / f"{base_name}.docx"

    doc.save(str(output_path))
    print(f"DOCX resume rendered successfully to: {output_path}")

    try:
        shutil.copy2(str(output_path), str(downloads_path))
        print(f"Copy saved to Downloads folder: {downloads_path}")
    except Exception as e:
        print(f"Warning: Could not copy to Downloads folder: {e}")


if __name__ == "__main__":
    arg = sys.argv[1] if len(sys.argv) > 1 else None
    render_docx(arg)
